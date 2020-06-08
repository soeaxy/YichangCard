# -*- coding: cp936 -*-
import string
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

import arcpy
from ExportJPG import ExportToJPGFile

# �ж���ʽ�Ƿ����
def styleExists(document, style_name):
    stylename_list = [st.name for st in document.styles]
    if style_name in stylename_list:
        return True
    else:
        return False

# �����ʽ���ĵ�
def addStyleToDocument(document, style_name):
    if not styleExists(document,style_name):
        title_style = document.styles.add_style(style_name, WD_ALIGN_PARAGRAPH)
        title_style.font.size = Pt(16)
        title_style.font.bold=True

# ����һ�����
def CreateTable(document,row, rsPicPath=None):
    rows,cols = 12,25
    table = document.add_table(rows=rows,cols=cols,style = 'Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # ����������
    table.style.font.size=Pt(11) # �ֺ�
    for i in range(rows):
        table.rows[i].height = Cm(0.6)

    
    # ���ڱ��
    table.cell(0,0).merge(table.cell(0,2)).text=(u"������")
    table.cell(0,3).merge(table.cell(0,6)).text=(row.getValue(u"������"))

    # ��Ŀ����
    table.cell(0,7).merge(table.cell(0,11)).text=(u"����Ŀ����")
    table.cell(0,12).merge(table.cell(0,24)).text=(u'й̲��1:5��ˮ�ĵ��ʻ��������ۺ�ң�н���')

    # ��������
    table.cell(1,0).merge(table.cell(2,2)).text=(u"��������")
    table.cell(1,3).merge(table.cell(2,8)).text=(u"�ϲ�")
    table.cell(1,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #����
    table.cell(1,9).merge(table.cell(2,10)).text=(u"����")
    table.cell(1,9).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Y
    table.cell(2,11).merge(table.cell(2,12)).text=(u"Y")
    table.cell(2,13).merge(table.cell(2,16)).text=("{0:.2f}".format(row.getValue(u"Y����")))

    # X
    table.cell(1,11).merge(table.cell(1,12)).text=(u"X")
    table.cell(1,13).merge(table.cell(1,16)).text=("{0:.2f}".format(row.getValue(u"X����")))

    # N
    table.cell(1,17).merge(table.cell(1,18)).text=(u"N")
    table.cell(1,19).merge(table.cell(1,24)).text=(row.getValue(u"γ��"))

    # E
    table.cell(2,17).merge(table.cell(2,18)).text=(u"E")
    table.cell(2,19).merge(table.cell(2,24)).text=(row.getValue(u"����"))

    # ����λ��
    table.cell(3,0).merge(table.cell(3,2)).text=(u'����λ��')
    table.cell(3,3).merge(table.cell(3,24)).text=(row.getValue(u'Problem'))
     
    # Ӱ���ò
    table.cell(4,0).merge(table.cell(4,2)).text=(u'Ӱ���ò')
    table.cell(4,3).merge(table.cell(4,24)).text=(row.getValue(u'Problem'))

    # ң��Ӱ��
    table.cell(5,0).merge(table.cell(5,2)).text=(u"ң��Ӱ��")
    table.cell(5,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_polygonImage = table.cell(5,3).merge(table.cell(5,24))
    pic_p = c_polygonImage.paragraphs[0]
    pic_r = pic_p.add_run()
    pic_r.add_picture(rsPicPath,width=Inches(3))
    pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ұ����֤
    table.cell(6,0).merge(table.cell(6,2)).text=(u"Ұ����֤")
    table.cell(6,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(6,3).merge(table.cell(6,13))
    table.cell(6,14).merge(table.cell(6,24))
    table.rows[6].height=Inches(3)

    # �ۺϽ���
    table.cell(7,0).merge(table.cell(9,2)).text=(u'�ۺϽ���')
    table.cell(7,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(7,3).merge(table.cell(9,24)).text=(row.getValue(u'����'))

    # ��Ŀ��Ϣ
    table.cell(10,0).merge(table.cell(10,2)).text=(u"��Ŀ����")
    table.cell(10,3).merge(table.cell(10,15)).text=(u'�˲�����Դ����������������')
    table.cell(10,16).merge(table.cell(10,18)).text=(u"�е���λ")
    table.cell(10,19).merge(table.cell(10,24)).text=(u'')

    # ��Ա��Ϣ
    table.cell(11,0).merge(table.cell(11,2)).text=(u"������Ա")
    table.cell(11,3).merge(table.cell(11,8)).text=(u'')
    table.cell(11,9).merge(table.cell(11,12)).text=(u"����ʱ��")
    table.cell(11,13).merge(table.cell(11,18)).text=('')
    table.cell(11,19).merge(table.cell(11,21)).text=(u'���')
    table.cell(11,22).merge(table.cell(11,24)).text=('')

#----------ArcTooolboxģʽ----------
# docx_template_file = arcpy.GetParameterAsText(0)
# table_title = arcpy.GetParameterAsText(1)
# images_output_dir = arcpy.GetParameterAsText(2)
# para_fc_count = arcpy.GetParameterAsText(3)
# mxd = arcpy.mapping.MapDocument("CURRENT")

# # #----------��������ģʽ----------
docx_template_file = "XietanFaultResult.docx"
table_title = u""
images_output_dir = r"."
# # mxd�ļ��мǣ�������������������ᱨIO��
mxd = arcpy.mapping.MapDocument(r"XietanFaultForExport.mxd")

lyr = arcpy.mapping.ListLayers(mxd)[0]
df = arcpy.mapping.ListDataFrames(mxd)[0]

fields = arcpy.ListFields(lyr.dataSource)
fieldnames = [f.name for f in fields]

document = Document(docx_template_file)
cursor = arcpy.SearchCursor(lyr,fieldnames)

# ��para_fc_countΪ0�����������Ҫ��
# ��para_fc_count��Ϊ0�������para_fc_count��Ҫ��
# fc_countΪ���������docx�ļ���Ҫ����Ŀ
para_fc_count=0
fc_count=0
if int(para_fc_count)==0:
    fc_count = int(arcpy.GetCount_management(lyr.dataSource).getOutput(0))
else:
    fc_count = int(para_fc_count)

# ���ý�����
arcpy.SetProgressor("step", u"�������ɽ��뿨Ƭ...", 0, fc_count, 1)

ii=0
for row in cursor:
    if ii<fc_count:
        document.add_paragraph(table_title, "MyStyle")
        filename = row.getValue(u"������")
        picPath= ExportToJPGFile(mxd,df,lyr,ii,images_output_dir,filename)
        CreateTable(document, row, picPath)
        arcpy.AddMessage(u"��ɵ� {0} �� ��Ƭ...".format(ii+1))
        document.add_page_break()
        # ���½�����λ��
        arcpy.SetProgressorPosition()
    else:
        break
    ii = ii + 1
document.save(docx_template_file)

